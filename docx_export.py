"""Generate a Word (.docx) version of the SFBC 2025 Budget Report."""

import io
import os
from datetime import date

from config import ANALYSIS_END
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
ACCENT   = RGBColor(0x29, 0x62, 0x9B)   # dark blue
SUBHEAD  = RGBColor(0x3C, 0x64, 0xA0)
HDR_BG   = RGBColor(0x29, 0x62, 0x9B)
HDR_FG   = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT  = RGBColor(0xF0, 0xF5, 0xFC)
NOTE_BG  = RGBColor(0xFF, 0xFC, 0xE6)
POS_CLR  = RGBColor(0x1E, 0x82, 0x4C)
NEG_CLR  = RGBColor(0xB4, 0x1E, 0x1E)
DARK     = RGBColor(0x1E, 0x1E, 0x1E)
GRAY     = RGBColor(0xB4, 0xB4, 0xB4)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _cell_text(cell, text: str, bold: bool = False,
               color: RGBColor | None = None,
               align: str = "left",
               font_size: int = 9):
    para = cell.paragraphs[0]
    para.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level <= 1 else SUBHEAD
    return p


def _para(doc: Document, text: str, italic: bool = False,
          color: RGBColor | None = None, size: int = 10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _kv_para(doc: Document, label: str, value: str,
             value_bold: bool = False, value_color: RGBColor | None = None):
    p = doc.add_paragraph()
    r_lbl = p.add_run(f"{label}  ")
    r_lbl.font.size = Pt(10)
    r_val = p.add_run(value)
    r_val.bold = value_bold
    r_val.font.size = Pt(10)
    if value_color:
        r_val.font.color.rgb = value_color
    return p


def _bullet(doc: Document, text: str, size: int = 9):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _add_table(doc: Document, headers: list, rows: list,
               col_widths_in: list, alignments: list | None = None):
    """Add a styled table to the document."""
    alignments = alignments or (["left"] * len(headers))
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (h, w_in) in enumerate(zip(headers, col_widths_in)):
        cell = hdr_cells[i]
        _set_cell_bg(cell, _rgb_hex(HDR_BG))
        _cell_text(cell, h, bold=True, color=HDR_FG,
                   align=alignments[i], font_size=9)
        cell.width = Inches(w_in)

    # Body rows
    for r_idx, row in enumerate(rows):
        is_total = str(row[0]).upper().replace("*", "").strip().startswith("TOTAL")
        row_cells = table.rows[r_idx + 1].cells
        for i, (val, w_in) in enumerate(zip(row, col_widths_in)):
            cell = row_cells[i]
            if is_total:
                _set_cell_bg(cell, _rgb_hex(HDR_BG))
                _cell_text(cell, val, bold=True, color=HDR_FG,
                           align=alignments[i], font_size=9)
            elif r_idx % 2 == 0:
                _set_cell_bg(cell, _rgb_hex(ROW_ALT))
                _cell_text(cell, val, align=alignments[i], font_size=9)
            else:
                _cell_text(cell, val, align=alignments[i], font_size=9)
            cell.width = Inches(w_in)

    doc.add_paragraph()   # spacing after table
    return table


def _add_image(doc: Document, path: str, caption: str, width_in: float = 6.0):
    """Insert a PNG image with a caption."""
    if not os.path.exists(path):
        return
    p = doc.add_picture(path, width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[attr-defined]
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main export function
# ---------------------------------------------------------------------------

def export_docx(
    report_data: dict,
    fidelity_data: dict,
    investment_data: dict,
    chart_dir: str,
    out_path: str,
) -> None:
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    r   = report_data
    f   = fidelity_data
    inv = investment_data or {}
    inv_year = ANALYSIS_END.year
    pl  = r.get("period_label", str(inv_year))
    doc = Document()
    generated = date.today().strftime("%B %d, %Y")

    # ---- Title -------------------------------------------------------------
    title = doc.add_heading(f"SFBC {pl} Budget Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT

    sub = doc.add_paragraph("San Francisco Bay Chapter")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY

    gen = doc.add_paragraph(f"Generated: {generated}")
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in gen.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY

    doc.add_paragraph()

    # ---- Summary -----------------------------------------------------------
    _heading(doc, "Summary", 1)
    total_income = r["total_membership"] + r["total_donations"]
    total_out    = r["total_expenses"] + r["total_paypal_fees"]
    net          = r["net_income"]
    net_color    = POS_CLR if net >= 0 else NEG_CLR

    summary_rows = [
        ["Membership Income",  f"${r['total_membership']:,.2f}"],
        ["Donations",          f"${r['total_donations']:,.2f}"],
        ["Total Income",       f"${total_income:,.2f}"],
        ["Expenses",           f"${r['total_expenses']:,.2f}"],
        ["PayPal Fees",        f"${r['total_paypal_fees']:,.2f}"],
        ["Total Outflows",     f"${total_out:,.2f}"],
        ["NET INCOME",         f"${net:,.2f}"],
    ]
    _add_table(doc, ["Metric", "Amount"], summary_rows,
               [3.5, 1.5], ["left", "right"])
    _add_image(doc, os.path.join(chart_dir, "chart_income_vs_expenses.png"),
               "Figure 1 - Income vs Expenses", width_in=5.5)

    # ---- Income Breakdown --------------------------------------------------
    _heading(doc, "Income Breakdown", 1)
    _heading(doc, "Monthly Income", 2)
    im = r["income_monthly"].copy()
    im["Month"] = im["Month"].str.replace("**", "", regex=False)
    _add_table(doc, ["Month", "Membership", "Donations", "Total"],
               im.values.tolist(),
               [1.5, 1.3, 1.3, 1.4], ["left", "right", "right", "right"])
    _add_image(doc, os.path.join(chart_dir, "chart_monthly_cashflow.png"),
               "Figure 2 - Monthly Cash Flow", width_in=6.5)

    # ---- Expense Breakdown -------------------------------------------------
    _heading(doc, "Expense Breakdown", 1)
    _heading(doc, "Monthly Expenses", 2)
    em = r["expense_monthly"].copy()
    em["Month"] = em["Month"].str.replace("**", "", regex=False)
    _add_table(doc, ["Month", "Expenses"], em.values.tolist(),
               [3.5, 1.5], ["left", "right"])
    _add_image(doc, os.path.join(chart_dir, "chart_expense_breakdown.png"),
               "Figure 3 - Expense Breakdown by Category", width_in=4.5)

    _heading(doc, "Major Expenses (>$50) with Notes", 2)
    if not r["major_expenses"].empty:
        me = r["major_expenses"]
        _add_table(doc, ["Date", "Description", "Amount", "Note"],
                   me.values.tolist(),
                   [0.9, 2.2, 0.9, 2.0], ["left", "left", "right", "left"])
    else:
        _para(doc, "No major expenses found.", italic=True)

    _heading(doc, "Zelle Payments Out", 2)
    if not r["zelle_out"].empty:
        zo = r["zelle_out"]
        _add_table(doc, ["Date", "Description", "Amount", "Note"],
                   zo.values.tolist(),
                   [0.9, 2.2, 0.9, 2.0], ["left", "left", "right", "left"])
    else:
        _para(
            doc,
            f"No outgoing Zelle payments between {r['analysis_start']} and {r['analysis_end']}.",
            italic=True,
        )

    _heading(doc, "All Itemized Expenses", 2)
    ei = r["expense_items"]
    _add_table(doc, ["Date", "Description", "Amount"],
               ei.values.tolist(),
               [0.9, 4.2, 0.9], ["left", "left", "right"])

    # ---- PayPal Summary ----------------------------------------------------
    _heading(doc, "PayPal Summary", 1)
    if r["paypal_has_2025_data"]:
        paypal_rows = [
            ["Gross Received", f"${r['paypal_gross']:,.2f}"],
            ["Fees",           f"${r['total_paypal_fees']:,.2f}"],
            ["Net",            f"${r['paypal_net']:,.2f}"],
        ]
        _add_table(doc, ["Item", "Amount"], paypal_rows,
                   [3.5, 1.5], ["left", "right"])
    else:
        _para(doc,
              f"No completed PayPal income for {r['analysis_start']}–{r['analysis_end']}, "
              "or amounts net to zero. Confirm paypal.CSV has Completed rows in that range.",
              italic=True)

    # ---- Fidelity Investments ----------------------------------------------
    _heading(doc, "Fidelity Investments", 1)
    if f and f["total_value"] > 0:
        _kv_para(doc, "Account:", f["account_number"])
        _kv_para(doc, "Current Portfolio Value:",
                 f"${f['total_value']:,.2f}", value_bold=True)
        doc.add_paragraph()

        if inv:
            _heading(doc, f"{inv_year} Annual Investment Performance", 2)
            ret     = inv["total_return"]
            ret_pct = inv["return_pct"]
            v_color = POS_CLR if ret >= 0 else NEG_CLR
            perf_rows = [
                [f"Beginning Balance (Jan 1, {inv_year})",  f"${inv['beginning_balance']:,.2f}"],
                ["Market Change",                    f"${inv['market_change']:+,.2f}"],
                ["Dividends",                        f"${inv['dividends']:,.2f}"],
                [f"Ending Balance (Dec 31, {inv_year})",    f"${inv['ending_balance']:,.2f}"],
                ["Total Return ($)",                 f"${ret:+,.2f}"],
                ["Total Return (%)",                 f"{ret_pct:+.2f}%"],
            ]
            _add_table(doc, ["Metric", "Value"], perf_rows,
                       [3.5, 1.5], ["left", "right"])

        _heading(doc, "Current Holdings", 2)
        holdings = f["holdings"].copy()
        holdings["current_value"]          = holdings["current_value"].map(lambda x: f"${x:,.2f}")
        holdings["last_price"]             = holdings["last_price"].map(
            lambda x: f"${x:,.2f}" if x > 0 else "-")
        holdings["total_gain_loss_dollar"] = holdings["total_gain_loss_dollar"].map(
            lambda x: f"+${x:,.2f}" if x >= 0 else f"-${abs(x):,.2f}")
        holdings["total_gain_loss_pct"]    = holdings["total_gain_loss_pct"].map(
            lambda x: f"{x:+.2f}%")
        holdings["percent_of_account"]     = holdings["percent_of_account"].map(
            lambda x: f"{x:.2f}%")
        fid_rows = holdings[[
            "symbol", "description", "last_price", "current_value",
            "total_gain_loss_dollar", "total_gain_loss_pct", "percent_of_account"
        ]].values.tolist()
        _add_table(
            doc,
            ["Symbol", "Description", "Price", "Value", "G/L $", "G/L %", "Alloc"],
            fid_rows,
            [0.6, 2.2, 0.7, 0.8, 0.9, 0.7, 0.6],
            ["left", "left", "right", "right", "right", "right", "right"],
        )
    else:
        _para(doc, "Fidelity data not available.", italic=True)

    # ---- Notes -------------------------------------------------------------
    _heading(doc, "Notes", 1)
    notes = [
        "Internal transfers excluded: PayPal <-> Chase ACH transfers are detected "
        "and excluded from all totals to avoid double-counting.",
        f"Date range: posting dates from {r['analysis_start']} through {r['analysis_end']} (inclusive).",
        "PayPal fees: Tracked separately and excluded from income.",
        "Donations: credits whose description matches donation-related keywords; other income is membership.",
        "Major expenses: Expenses >= $50 are annotated in the Major Expenses section.",
        "Name redaction: Personal names in transaction descriptions have been anonymized.",
        "Fidelity: Portfolio positions are as of the CSV export date; not mixed into operating figures.",
    ]
    for note in notes:
        _bullet(doc, note)

    if not r["transfers"].empty:
        _heading(doc, "Internal Transfers (Excluded)", 2)
        tf = r["transfers"]
        _add_table(doc, ["Date", "Description", "Amount", "Source"],
                   tf.values.tolist(),
                   [0.9, 4.4, 0.9, 0.8], ["left", "left", "right", "left"])

    doc.save(out_path)
    print(f"  DOCX written to {out_path}")
